"""DOCX text extraction, including paragraph text and tables."""

import docx


def extract_text_from_docx(file_path: str) -> str:
    """Extract paragraph and table text from a Word document."""
    document = docx.Document(file_path)
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts)
